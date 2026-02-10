"""
文档解析服务

支持多种文档格式的文本提取：
- PDF (.pdf)
- Word (.docx, .doc)
- 纯文本 (.txt)
- Markdown (.md)
"""

import io
import re
from pathlib import Path
from typing import Optional, Dict, Any, List
from loguru import logger


class DocumentParser:
    """文档解析器"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md'}
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查可用的解析库"""
        self.has_pypdf = False
        self.has_docx = False
        
        try:
            import pypdf
            self.has_pypdf = True
        except ImportError:
            logger.warning("pypdf 未安装，PDF解析功能受限")
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            logger.warning("python-docx 未安装，Word文档解析功能受限")
    
    async def parse_file(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        file_name: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        解析文档文件
        
        Args:
            file_path: 文件路径
            file_content: 文件二进制内容
            file_name: 文件名（用于判断类型）
            
        Returns:
            解析结果字典
        """
        if file_path:
            path = Path(file_path)
            file_name = path.name
            with open(path, 'rb') as f:
                file_content = f.read()
        
        if not file_content or not file_name:
            return {"error": "未提供有效的文件内容", "text": ""}
        
        ext = Path(file_name).suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            return {
                "error": f"不支持的文件格式: {ext}",
                "text": "",
                "supported_formats": list(self.SUPPORTED_EXTENSIONS)
            }
        
        try:
            if ext == '.pdf':
                text = await self._parse_pdf(file_content)
            elif ext in ['.docx', '.doc']:
                text = await self._parse_docx(file_content)
            elif ext in ['.txt', '.md']:
                text = self._parse_text(file_content)
            else:
                text = ""
            
            # 提取结构化信息
            structure = self._extract_structure(text)
            
            return {
                "success": True,
                "text": text,
                "char_count": len(text),
                "word_count": len(text.split()),
                "structure": structure,
                "file_name": file_name,
                "file_type": ext,
            }
            
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            return {
                "error": str(e),
                "text": "",
                "file_name": file_name,
            }
    
    async def _parse_pdf(self, content: bytes) -> str:
        """解析 PDF 文件"""
        if not self.has_pypdf:
            # 降级：尝试使用 pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    return "\n\n".join(text_parts)
            except ImportError:
                raise ImportError("需要安装 pypdf 或 pdfplumber 来解析PDF文件")
        
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    async def _parse_docx(self, content: bytes) -> str:
        """解析 Word 文档"""
        if not self.has_docx:
            raise ImportError("需要安装 python-docx 来解析Word文档")
        
        import docx
        doc = docx.Document(io.BytesIO(content))
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _parse_text(self, content: bytes) -> str:
        """解析纯文本文件"""
        # 尝试多种编码
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode('utf-8', errors='ignore')
    
    def _extract_structure(self, text: str) -> Dict[str, Any]:
        """提取文档结构信息"""
        structure = {
            "sections": [],
            "has_signature": False,
            "has_date": False,
            "has_parties": False,
        }
        
        # 检测章节标题
        section_patterns = [
            r'第[一二三四五六七八九十百]+[条章节]',  # 第一条, 第十二条
            r'[一二三四五六七八九十]+、',           # 一、
            r'\d+\.\s*\w+',                       # 1. 标题
            r'[（\(][\d一二三四五六七八九十]+[）\)]', # (1) 或 (一)
            r'第[0-9]+[条章节]',                   # 第1条
            r'附则',                              # 附则
            r'前言',                              # 前言
        ]
        
        for pattern in section_patterns:
            matches = re.findall(pattern, text)
            if matches:
                structure["sections"].extend(matches[:20])  # 限制数量
        
        # 检测签名区域
        signature_keywords = ['签字', '签章', '盖章', '签名', '法定代表人', '授权代表']
        structure["has_signature"] = any(kw in text for kw in signature_keywords)
        
        # 检测日期
        date_pattern = r'\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{2}-\d{2}|\d{4}/\d{2}/\d{2}'
        structure["has_date"] = bool(re.search(date_pattern, text))
        
        # 检测合同主体
        party_keywords = ['甲方', '乙方', '丙方', '出租方', '承租方', '买方', '卖方', '委托方', '受托方']
        structure["has_parties"] = any(kw in text for kw in party_keywords)
        
        return structure


class ContractTextAnalyzer:
    """合同文本分析器"""
    
    # 常见合同类型关键词
    CONTRACT_TYPE_KEYWORDS = {
        "租赁合同": ["租赁", "出租", "承租", "租金", "房屋", "场地"],
        "劳动合同": ["劳动", "员工", "工资", "薪酬", "社保", "劳动者"],
        "买卖合同": ["买卖", "购买", "销售", "货物", "商品", "采购"],
        "服务合同": ["服务", "咨询", "技术服务", "外包", "顾问"],
        "借款合同": ["借款", "贷款", "利息", "还款", "担保"],
        "股权转让": ["股权", "股份", "转让", "股东", "投资"],
        "保密协议": ["保密", "商业秘密", "机密信息", "NDA"],
        "合作协议": ["合作", "联营", "合资", "共同开发"],
    }
    
    # 高风险条款关键词
    HIGH_RISK_KEYWORDS = [
        "无条件", "不可撤销", "放弃追索", "全部责任", "无限责任",
        "自动续约", "单方解除", "排他性", "竞业禁止", "连带责任",
        "不可抗力免责", "争议仲裁", "损害赔偿上限",
    ]
    
    # 需要特别关注的条款类型
    IMPORTANT_CLAUSE_TYPES = [
        "付款条款", "违约责任", "保密条款", "知识产权", 
        "争议解决", "合同解除", "不可抗力", "损害赔偿",
    ]
    
    def analyze_contract_type(self, text: str) -> str:
        """分析合同类型"""
        scores = {}
        for contract_type, keywords in self.CONTRACT_TYPE_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in text)
            if score > 0:
                scores[contract_type] = score
        
        if scores:
            return max(scores, key=scores.get)
        return "通用合同"
    
    def extract_key_info(self, text: str) -> Dict[str, Any]:
        """提取合同关键信息"""
        info = {
            "parties": self._extract_parties(text),
            "amount": self._extract_amount(text),
            "dates": self._extract_dates(text),
            "high_risk_terms": self._find_high_risk_terms(text),
        }
        return info
    
    def _extract_parties(self, text: str) -> List[str]:
        """提取合同主体"""
        parties = []
        
        # 匹配甲方/乙方模式
        patterns = [
            r'甲方[：:]\s*([^，,。\n]+)',
            r'乙方[：:]\s*([^，,。\n]+)',
            r'出租方[：:]\s*([^，,。\n]+)',
            r'承租方[：:]\s*([^，,。\n]+)',
            r'买方[：:]\s*([^，,。\n]+)',
            r'卖方[：:]\s*([^，,。\n]+)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            parties.extend(matches)
        
        return list(set(parties))[:4]  # 去重并限制数量
    
    def _extract_amount(self, text: str) -> Optional[str]:
        """提取合同金额"""
        patterns = [
            r'合同金额[：:为]?\s*[人民币RMB￥¥]?\s*([\d,，.]+)\s*[元万亿]',
            r'总价[：:为]?\s*[人民币RMB￥¥]?\s*([\d,，.]+)\s*[元万亿]',
            r'[人民币RMB￥¥]\s*([\d,，.]+)\s*[元万亿]',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return None
    
    def _extract_dates(self, text: str) -> List[str]:
        """提取日期"""
        pattern = r'\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{2}-\d{2}|\d{4}/\d{2}/\d{2}'
        dates = re.findall(pattern, text)
        return list(set(dates))[:5]
    
    def _find_high_risk_terms(self, text: str) -> List[str]:
        """查找高风险条款"""
        found = []
        for term in self.HIGH_RISK_KEYWORDS:
            if term in text:
                found.append(term)
        return found


# 创建全局实例
document_parser = DocumentParser()
contract_analyzer = ContractTextAnalyzer()


async def parse_contract_document(
    file_path: Optional[str] = None,
    file_content: Optional[bytes] = None,
    file_name: Optional[str] = None,
) -> Dict[str, Any]:
    """
    解析合同文档的便捷函数
    
    Returns:
        包含文本内容和分析结果的字典
    """
    # 解析文档
    result = await document_parser.parse_file(
        file_path=file_path,
        file_content=file_content,
        file_name=file_name,
    )
    
    if result.get("error"):
        return result
    
    text = result.get("text", "")
    
    # 分析合同类型和关键信息
    result["contract_type"] = contract_analyzer.analyze_contract_type(text)
    result["key_info"] = contract_analyzer.extract_key_info(text)
    
    return result
